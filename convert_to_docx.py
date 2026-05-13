"""Convert CLAUDE.md to a formatted Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD_PATH  = r'C:\Users\RyanCWalsh\.claude\ga-automation\CLAUDE.md'
OUT_PATH = r'C:\Users\RyanCWalsh\.claude\ga-automation\GA_Automation_Pipeline_Overview.docx'

with open(MD_PATH, encoding='utf-8') as f:
    lines = f.readlines()

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ── Helpers ───────────────────────────────────────────────────────────

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_code_block(doc, lines_list):
    for ln in lines_list:
        p = doc.add_paragraph(ln if ln else ' ')
        p.style = 'No Spacing'
        if not p.runs:
            p.add_run(ln)
        run = p.runs[0]
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        for side in ('top', 'left', 'bottom', 'right'):
            bd = OxmlElement(f'w:{side}')
            bd.set(qn('w:val'),   'single')
            bd.set(qn('w:sz'),    '4')
            bd.set(qn('w:space'), '4')
            bd.set(qn('w:color'), 'C0C0C0')
            pBdr.append(bd)
        pPr.append(pBdr)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F5F5F5')
        pPr.append(shd)


def flush_table(doc, tbl_rows):
    if not tbl_rows:
        return
    data = [r for r in tbl_rows
            if not all(c.strip().replace('-', '').replace('|', '') == '' for c in r)]
    if not data:
        return
    cols = max(len(r) for r in data)
    tbl  = doc.add_table(rows=0, cols=cols)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(data):
        tr = tbl.add_row()
        for ci in range(cols):
            cell_txt = row[ci].strip() if ci < len(row) else ''
            cell = tr.cells[ci]
            cell.text = cell_txt
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    if ri == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if ri == 0:
                shade_cell(cell, '1F4E78')


def parse_inline(text):
    """Return list of (chunk, bold, code) from markdown inline markup."""
    parts   = []
    pattern = re.compile(r'(\*\*(.+?)\*\*|`([^`]+)`)')
    last    = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False, False))
        if m.group(0).startswith('**'):
            parts.append((m.group(2), True, False))
        else:
            parts.append((m.group(3), False, True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False, False))
    return parts


def add_para_with_inline(doc, text, style='Normal', indent=0):
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)   # strip links
    p    = doc.add_paragraph(style=style)
    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.25)
    for chunk, bold, is_code in parse_inline(text):
        run = p.add_run(chunk)
        run.font.size = Pt(10)
        if bold:
            run.bold = True
        if is_code:
            run.font.name  = 'Courier New'
            run.font.size  = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    return p


# ── Main pass ─────────────────────────────────────────────────────────
in_code  = False
code_buf = []
in_table = False
tbl_rows = []

for raw_line in lines:
    line = raw_line.rstrip('\n')

    # ── Code block ──
    if line.strip().startswith('```'):
        if in_code:
            add_code_block(doc, code_buf)
            code_buf.clear()
            in_code = False
        else:
            if in_table:
                flush_table(doc, tbl_rows)
                tbl_rows.clear()
                in_table = False
            in_code = True
        continue
    if in_code:
        code_buf.append(line)
        continue

    # ── Table row ──
    if line.strip().startswith('|'):
        in_table = True
        cells = line.split('|')[1:-1]
        tbl_rows.append(cells)
        continue
    elif in_table:
        flush_table(doc, tbl_rows)
        tbl_rows.clear()
        in_table = False

    # ── Headings ──
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        level = min(len(m.group(1)), 4)
        doc.add_heading(m.group(2).strip(), level=level)
        continue

    # ── Horizontal rule ──
    if re.match(r'^---+\s*$', line):
        p   = doc.add_paragraph('', style='No Spacing')
        run = p.add_run('─' * 90)
        run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        run.font.size      = Pt(8)
        continue

    # ── Blank line ──
    if not line.strip():
        doc.add_paragraph('', style='No Spacing')
        continue

    # ── Bullet ──
    m = re.match(r'^(\s*)[-*]\s+(.*)', line)
    if m:
        depth = len(m.group(1)) // 2
        add_para_with_inline(doc, m.group(2), style='List Bullet', indent=depth)
        continue

    # ── Numbered list ──
    m = re.match(r'^(\s*)\d+\.\s+(.*)', line)
    if m:
        depth = len(m.group(1)) // 2
        add_para_with_inline(doc, m.group(2), style='List Number', indent=depth)
        continue

    # ── Normal paragraph ──
    add_para_with_inline(doc, line)

# Flush any trailing blocks
if in_code:
    add_code_block(doc, code_buf)
if in_table:
    flush_table(doc, tbl_rows)

doc.save(OUT_PATH)
print('Saved:', OUT_PATH)
